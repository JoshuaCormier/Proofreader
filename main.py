import re
import os
import glob
from docx import Document


class SmokeControlConsoleAuditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.current_year = 2026
        self.issues_found = 0

        self.project_profile = {
            "type": None,
            "city": None,
            "state": None
        }

        # Hardened regex patterns mapping technical anomalies to console feedback directives
        self.static_rules = [
            (r'\b(JJune|JJannuary|FFebruary|MMarch|AApril|MMay)\b',
             "Spelling Error: Double letter typo detected in month naming convention."),
            (r'\b([a-zA-Z]+)\s+\1\b', "Grammar Error: Duplicate consecutive words detected."),
            (r'Table\s{2,}\d+|Figure\s{2,}\d+',
             "Formatting Error: Excessive whitespace spacing inside caption identifier label."),
            (r'Error!\s*Reference\s*source\s*not\s*found|Error!\s*$',
             "Broken Reference: Missing equation reference template anchor or broken Word field code."),
            (r'\bScenario\s+X\b', "Placeholder Error: Uncleared scenario variable placeholder string left in text.")
        ]

    def extract_project_profile(self):
        """Isolates the document text to establish the dynamic tracking profile template."""
        all_body_text = " ".join([p.text for p in self.doc.paragraphs if p.text])

        types_pool = ["arena", "stadium", "convention center", "atrium", "high-rise"]
        for t in types_pool:
            if re.search(rf'\b{t}\b', all_body_text, re.IGNORECASE):
                self.project_profile["type"] = t
                break

        location_match = re.search(r'\b([A-Z][a-zA-Z\s]+),\s([A-Z]{2})\b', all_body_text)
        if location_match:
            self.project_profile["city"] = location_match.group(1).strip()
            self.project_profile["state"] = location_match.group(2).strip()

    def count_expected_figures(self, text):
        """Parses text to calculate exactly how many figures the author claims to reference."""
        range_match = re.search(r'Figures?\s+(\d+)\s*(?:through|-)\s*(\d+)', text, re.IGNORECASE)
        if range_match:
            start = int(range_match.group(1))
            end = int(range_match.group(2))
            return (end - start) + 1

        list_match = re.search(r'Figures?\s+(\d+(?:\s*,\s*\d+)*\s*(?:and|&)\s*\d+)', text, re.IGNORECASE)
        if list_match:
            digits = re.findall(r'\d+', list_match.group(1))
            return len(digits)

        single_match = re.search(r'\bFigures?\s+(\d+)\b', text, re.IGNORECASE)
        if single_match:
            return 1

        return 0

    def check_xml_graphic_elements(self, paragraph):
        """Counts the raw inline shapes or drawing objects inside a paragraph's XML tree."""
        p_element = paragraph._p
        drawings = p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        shapes = p_element.findall('.//{urn:schemas-microsoft-com:vml}shape')
        return len(drawings) + len(shapes)

    def print_to_console(self, layer, index, error_type, snippet, message):
        """Formats and outputs clear, itemized issue text block indicators directly to PyCharm."""
        self.issues_found += 1
        clean_snippet = snippet.strip().replace('\n', ' ')
        if len(clean_snippet) > 60:
            clean_snippet = clean_snippet[:57] + "..."

        print(f"[{self.issues_found}] LOCATION: {layer} (Item/Line {index})")
        print(f"    CLASSIFICATION: {error_type}")
        print(f"    TARGET TEXT:    \"{clean_snippet}\"")
        print(f"    DIRECTIVE:      {message}")
        print("-" * 110)

    def inspect_text_segment(self, paragraph, text, layer_name, index):
        """Checks text blocks against syntax patterns, profiles, calendar dates, and figure counts."""
        if not text.strip():
            return

        # 1. Isolate and Explicitly Inspect Native Metadata Captions
        is_caption_style = paragraph.style.name.lower() == 'caption'
        is_caption_text = text.strip().lower().startswith(('figure', 'table')) and (
                    ':' in text or re.search(r'\d', text))

        if is_caption_style or is_caption_text:
            # Enforce unified spacing check on the caption string label
            if re.search(r'(?:Table|Figure)\s{2,}\d+', text):
                self.print_to_console(layer_name, index, "Caption Layout Error", text,
                                      "Excessive whitespace gaps identified within the label prefix syntax structure.")

        # 2. Evaluate Static Syntax and Broken Reference RegEx Arrays
        for pattern, error_msg in self.static_rules:
            if re.search(pattern, text, re.IGNORECASE):
                self.print_to_console(layer_name, index, "Syntax/Formatting Exception", text, error_msg)

        # 3. Evaluate Dynamic Figure Reference Concordance
        claimed_count = self.count_expected_figures(text)
        if claimed_count > 0:
            actual_embedded_images = self.check_xml_graphic_elements(paragraph)
            if actual_embedded_images > 0 and claimed_count != actual_embedded_images:
                msg = f"Figure Reference Mismatch: The statement text claims to reference {claimed_count} figures, but {actual_embedded_images} physical graphic element(s) were found attached nearby."
                self.print_to_console(layer_name, index, "Graphic Concordance Exception", text, msg)

        # 4. Evaluate Asset and Boundary Profile Configurations
        target_type = self.project_profile.get("type")
        target_city = self.project_profile.get("city")

        if target_type:
            conflicts = {"arena", "stadium", "convention center", "atrium", "high-rise"} - {target_type.lower()}
            for conflict in conflicts:
                if re.search(rf'\b{conflict}\b', text, re.IGNORECASE):
                    # Skip standard occurrences inside table of contents layouts
                    if "table of contents" in text.lower() or text.strip().lower().startswith('figure'):
                        continue
                    msg = f"Asset Profile Conflict: Document references variant asset classification '{conflict}', mismatching Project Description profile context '{target_type}'."
                    self.print_to_console(layer_name, index, "Profile Verification Failure", text, msg)

        if target_city and "oklahoma city" in target_city.lower():
            if "mccarran" in text.lower() or "las vegas" in text.lower():
                msg = "Template Boilerplate Contamination: Climate data calculations reference McCarran Airport (Las Vegas) metrics instead of Oklahoma City telemetry."
                self.print_to_console(layer_name, index, "Profile Verification Failure", text, msg)

        # 5. Evaluate Date Parameters against active year criteria
        date_pattern = r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
        written_date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+(\d{4})\b'

        for match in re.finditer(date_pattern, text):
            date_str = match.group(0)
            if not any(str(self.current_year) in date_str or f"/{str(self.current_year)[2:]}" in date_str for _ in [0]):
                msg = f"Temporal Lifespan Failure: Stale numerical date context '{date_str}' deviates from running verification year ({self.current_year})."
                self.print_to_console(layer_name, index, "Stale Timeline Parameter", text, msg)

        for match in re.finditer(written_date_pattern, text):
            year = match.group(1)
            if int(year) != self.current_year:
                msg = f"Temporal Lifespan Failure: Stale written calendar year context '{year}' deviates from running verification year ({self.current_year})."
                self.print_to_console(layer_name, index, "Stale Timeline Parameter", text, msg)

    def drive_scanning_sequence(self):
        """Sequentially triggers textual evaluation blocks across all document component layers."""
        # Layer 1: Body Text Flow Elements
        for idx, paragraph in enumerate(self.doc.paragraphs):
            self.inspect_text_segment(paragraph, paragraph.text, "Main Body Text Paragraph", idx + 1)

        # Layer 2: Tabular Cell Matrix Grids
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        self.inspect_text_segment(paragraph, paragraph.text,
                                                  f"Table {t_idx + 1} Grid (Row {r_idx + 1}, Cell {c_idx + 1})",
                                                  p_idx + 1)

        # Layer 3: Dynamic Section Header Boundaries
        for s_idx, section in enumerate(self.doc.sections):
            headers = [section.header, section.first_page_header, section.even_page_header]
            for header in headers:
                if header:
                    for p_idx, paragraph in enumerate(header.paragraphs):
                        self.inspect_text_segment(paragraph, paragraph.text, f"Section {s_idx + 1} Header Workspace",
                                                  p_idx + 1)

            # Layer 4: Dynamic Section Footer Boundaries
            footers = [section.footer, section.first_page_footer, section.even_page_footer]
            for footer in footers:
                if footer:
                    for p_idx, paragraph in enumerate(footer.paragraphs):
                        self.inspect_text_segment(paragraph, paragraph.text, f"Section {s_idx + 1} Footer Workspace",
                                                  p_idx + 1)

        # Layer 5: Drawing Shape Graphic Floating Text boxes
        body_element = self.doc.element.body
        for b_idx, tx_bx in enumerate(
                body_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')):
            for p_node in tx_bx.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                from docx.text.paragraph import Paragraph
                temp_paragraph_wrapper = Paragraph(p_node, self.doc)
                self.inspect_text_segment(temp_paragraph_wrapper, temp_paragraph_wrapper.text,
                                          "Cover Page Text Box / Drawing Shape", b_idx + 1)

    def execute_review(self):
        self.extract_project_profile()

        print("=" * 110)
        print("SMOKE CONTROL RATIONAL ANALYSIS COMPLIANCE CHECKLIST ENGINE")
        print("=" * 110)
        print(f"Established Project Profile: {str(self.project_profile.get('type')).upper()}")
        print(
            f"Target Design Jurisdiction:  {str(self.project_profile.get('city')).upper()}, {str(self.project_profile.get('state')).upper()}")
        print(f"Validation Lifecycle Year:   {self.current_year}")
        print("=" * 110 + "\n")

        self.drive_scanning_sequence()

        print(f"\nCOMPLIANCE AUDIT COMPLETE: {self.issues_found} technical verification anomalies identified.")


if __name__ == "__main__":
    root_workspace = os.path.dirname(os.path.abspath(__file__))
    file_matches = glob.glob(os.path.join(root_workspace, "*.docx"))
    active_reports = [f for f in file_matches if not os.path.basename(f).startswith("~$")]

    if not active_reports:
        print(f"Execution Failure: No valid report files found in directory: {root_workspace}")
    else:
        target_report_file = active_reports[0]
        print(f"Auditing Source Asset: {os.path.basename(target_report_file)}\n")

        engine = SmokeControlConsoleAuditor(target_report_file)
        engine.execute_review()